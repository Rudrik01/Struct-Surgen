
import numpy as np
from docx import Document

def generate_report(filename, span_length, bay_length, number_of_bays, dead_load, live_load, wind_load):
    doc = Document()
    doc.add_heading('Industrial Shed Load Calculation Report', 0)
    
    doc.add_paragraph(f"Span Length: {span_length} meters")
    doc.add_paragraph(f"Bay Length: {bay_length} meters")
    doc.add_paragraph(f"Number of Bays: {number_of_bays}")
    doc.add_paragraph(f"Height of Ridge: {height_of_ridge} meters")
    doc.add_paragraph(f"Height of Eaves: {height_of_eaves} meters")
    
    doc.add_heading('Load Calculations', level=1)
    doc.add_paragraph(f"Dead Load: {dead_load:.2f} kN")
    doc.add_paragraph(f"Live Load: {live_load:.2f} kN")
    doc.add_paragraph(f"Wind Load: {wind_load:.2f} kN")
    
    doc.save(filename)
def generate_staad_file(filename, span_length, bay_length, number_of_bays, dead_load, live_load, wind_load):
    with open(filename, 'w') as file:
        file.write("STAAD SPACE\n")
        file.write("START JOB INFORMATION\n")
        file.write(f"Span Length: {span_length}\n")
        file.write(f"Bay Length: {bay_length}\n")
        file.write(f"Number of Bays: {number_of_bays}\n")
        file.write("END JOB INFORMATION\n")
        file.write("* Dead Load\n")
        file.write(f"LOAD 1 LOADTYPE Dead  TITLE Dead Load\n")
        file.write(f"MEMBER LOAD\n")
        file.write(f"1 TO {number_of_bays} UNI GY -{dead_load}\n")
        
        file.write("* Live Load\n")
        file.write(f"LOAD 2 LOADTYPE Live  TITLE Live Load\n")
        file.write(f"MEMBER LOAD\n")
        file.write(f"1 TO {number_of_bays} UNI GY -{live_load}\n")
        
        file.write("* Wind Load\n")
        file.write(f"LOAD 3 LOADTYPE Wind  TITLE Wind Load\n")
        file.write(f"MEMBER LOAD\n")
        file.write(f"1 TO {number_of_bays} UNI GX {wind_load}\n")

        file.write("FINISH\n")

def calculate_dead_load(span_length, bay_length, number_of_bays, height_of_eaves):
   
    density = 25 
    dead_load = density * span_length * bay_length * number_of_bays * height_of_eaves
    return dead_load

def calculate_live_load(span_length, bay_length):

    live_load = 2.5  
    return live_load * span_length * bay_length

def calculate_wind_load(span_length, height_of_eaves, height_of_ridge):
    wind_pressure = 0.6 
    projected_area = span_length * (height_of_ridge + height_of_eaves) / 2
    wind_load = wind_pressure * projected_area
    return wind_load

def get_shed_parameters():
    span_length = float(input("Enter the length of span (in meters): "))
    bay_length = float(input("Enter the bay length (in meters): "))
    number_of_bays = int(input("Enter the number of bays: "))
    height_of_ridge = float(input("Enter the height of the ridge (in meters): "))
    height_of_eaves = float(input("Enter the height of the eaves (in meters): "))
    return span_length, bay_length, number_of_bays, height_of_ridge, height_of_eaves
def main():
    
    span_length, bay_length, number_of_bays, height_of_ridge, height_of_eaves = get_shed_parameters()
    
   
    dead_load = calculate_dead_load(span_length, bay_length, number_of_bays, height_of_eaves)
    live_load = calculate_live_load(span_length, bay_length)
    wind_load = calculate_wind_load(span_length, height_of_eaves, height_of_ridge)
    
    
    report_filename = "industrial_shed_report.docx"
    generate_report(report_filename, span_length, bay_length, number_of_bays, dead_load, live_load, wind_load)
    
 
    staad_filename = "industrial_shed_structure.std"
    generate_staad_file(staad_filename, span_length, bay_length, number_of_bays, dead_load, live_load, wind_load)
    
    print(f"Report generated: {report_filename}")
    print(f"STAAD.Pro file generated: {staad_filename}")

if __name__ == "__main__":
    main()
