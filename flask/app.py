from flask import Flask, render_template, request, send_file
from docx import Document
import os
import math
import xlrd
from docx2pdf import convert
app = Flask(__name__)

document = Document()

# Route for home page, where form input will be provided
@app.route('/')
def home():
    return render_template('index.html')

# Route to handle form submission and generate Word document
@app.route('/generate', methods=['POST'])
def generate():
    # Collecting input from the form
    S = float(request.form.get('span', 0))
    N = int(request.form.get('bays', 0))
    B = float(request.form.get('bay_length', 0))
    E = float(request.form.get('eves_height', 0))
    R = float(request.form.get('ridge_height', 0))
    LL = float(request.form.get('live_load', 0))
    Vb = int(request.form.get('vb', 0))
    K1 = float(request.form.get('k1', 0))
    cata = int(request.form.get('terrain_category', 0))
    K3 = float(request.form.get('k3', 0))
    WR = float(request.form.get('roof_weight', 0))
    WP = float(request.form.get('purline_weight', 0))
    SP = float(request.form.get('purline_spacing', 0))
    IWP = float(request.form.get('internal_wind_pressure', 0))

    document.add_heading("Load Calculation",level=0)
    L = float (B*N)
   
    S1=S/2
   
    gap = R-E
    A = round(math.degrees(math.tan(gap/S1)),2)

    print("LL With Access provide = 1.5 KN/m2 and without Access = 0.75 KN/m2")
   
    lr=math.sqrt((S1*S1)+(gap*gap))
    lr=round(lr,2)
    Ra=lr*B 
    Ra=round(Ra,2)
    LLF=Ra*LL
    LLF=round(LLF,2)
    LLR=LLF/lr



  
    #**********************************************************************************************************************************************
    #                                                                 DEAD LOAD
    #**********************************************************************************************************************************************
    document.add_paragraph('''Prelimnary Data''','heading 1')
    document.add_paragraph(f'''Span = {S} m
    Length = {L} m
    Eves Hight = {E} m
    Ridge Hight = {R} m
    ''')

    document.add_paragraph('''Prelimnary Calculation''','heading 1')
    document.add_paragraph(f'''
    Angle of Roof Truss 
            tan ø = {gap}/{S1}
                ø = {A}° ''')

    document.add_paragraph(f'''
    Length of Principle Rafter 
                    = √({gap}^2+{S1}^2 )
                    = {lr} m ''')

    document.add_paragraph(f'''
    Half Slope Area 
                    = {lr}*{B}
                    = {Ra}  ''')



    #**********************************************************************************************************************************************
    #                                                                 DEAD LOAD
    #**********************************************************************************************************************************************

    document.add_paragraph(''' Dead Load ''','heading 1')
  
    WRF = (WR*Ra)/lr
    WRF = round(WRF,2)
    print("Wright purline per meter length, ex-ISMC125=0.131 ")
  
    WPF = WP*SP
    WPF = round(WPF,2)
    TF = WRF + WPF 
    TF = round(TF,2)

    document.add_paragraph(f'''
    A.C.C. SHEET = {WR} KN/m2                        (IS-875, PART-1)
                            = {WR}*{Ra}
                            = {WRF} KN/m
    Wt. OF PURLIN = ({WP}*6)/{SP}                    (ISMC-125)
                                = {WPF} KN/m

    TOTAL D.L. = {WRF}+{WPF}
                        = {TF} KN/m ''')

    table=document.add_table(rows= N+1,cols=3)
    table.style = "Table Grid"  
    row = table.rows[0].cells
    row[0].text = "No."
    row[1].text = "Span"
    row[2].text = "Load"





    #**********************************************************************************************************************************************
    #                                                                LIVE LOAD
    #**********************************************************************************************************************************************
    document.add_paragraph(''' Live Load ''','heading 1')

    LLR=round(LLR,2)
    document.add_paragraph(f''' Roof Area = {lr}*{B} 
                        ={Ra} Sq.m''')
    document.add_paragraph(f''' Live Load = {Ra}*{LL} 
                        ={LLF} KN/m2''')
    document.add_paragraph(f''' Live Load on Rafter= {LLF}/{lr} 
                        ={LLR} KN/m2''')

    table=document.add_table(rows= N+1,cols=3)
    table.style = "Table Grid"  
    row = table.rows[0].cells
    row[0].text = "No."
    row[1].text = "Span"
    row[2].text = "Load"


    #**********************************************************************************************************************************************
    #                                                               WIND LOAD
    #**********************************************************************************************************************************************




    #                                                    design wind speed

    document.add_paragraph('''Design wind speed (Vz):''','heading 1')
    document.add_paragraph('''From Page-8
                Vz= Vb K1 K2 K3 m/sec''')
    document.add_paragraph(f'''Vb= Basic wind speed
                From Fig.1 or appendix-A of the code,
                            Vb= {Vb} m/sec''')
    document.add_paragraph(f'''K1= Risk cofficient 
                Table-1 page-11,
                            K1= {K1} ''')

    #**************  vale of k1 from table 2 ***************************** 
    # Give the location of the file
    loc = ("./table 2222.xls")

    wb = xlrd.open_workbook(loc)
    sheet = wb.sheet_by_index(0) # 

    start = [9,2]
    ending = [23,2]
    data = []
    for i in range(ending[0]-start[0]):
        j = i+start[0]
        data.append(float(sheet.cell_value(j, start[1])))
    i = 0
    for j in range(len(data)-1):
        if R> data[j] and R<data[j+1]:
            i = j
            break
    start0 = [start[0]+i,start[1]]
    ending0 = [start0[0]+1,start0[1]]

    # keys
    Az = (sheet.cell_value(start0[0], start0[1])) # 20
    Bz = (sheet.cell_value(ending0[0], ending0[1])) # 30
    start1 = [start[0]+i,start[1]+cata]
    ending1 = [start1[0]+1,start1[1]]

    # values
    A1 = (sheet.cell_value(start1[0], start1[1])) # 1.12
    B1 = (sheet.cell_value(ending1[0], ending1[1])) # 1.15

    U = (((R-Az)*(B1-A1))/(Bz-Az))+A1
    K2 = U     

    document.add_paragraph(f'''K2= terrain, height and structure size coefficient
                Table-2 page-12,
                            K2= {K2} ''')


    #********************************************************************                         
    document.add_paragraph(f'''K3= topography factor
                Cl 5.3.3.1 page-12,
                            K3= {K3} ''')
    Vz=Vb*K1*K2*K3
    document.add_paragraph(f'''Now,
                Vz= Vb K1 K2 K3
                    = {Vb}*{K1}*{K2}*{K3}
                    = {Vz} m/sec''')

    #                                    Design wind Presure 

    document.add_paragraph('''Design wind pressure (Pz):''')
    Pz= (0.6*Vz**2)/1000
    document.add_paragraph(f'''
                Pz= 0.6 Vz^2
                    = 0.6*{Vz}^2
                    = {Pz} Kn/m2 ''')

    #                                     Wind Load

    print('''Cl No. 6.2.3.2 Buildings with openings 5 to 20 percent of wall area 0.5
                                openings larger than 20 percent of the wall area 0.7''')

   
    IWPPLUS = IWP
    IWPMINUS = -IWP

    # Give the location of the file
    loc = ("./table 5555.xls")

    wb = xlrd.open_workbook(loc)
    sheet = wb.sheet_by_index(0) # 

    # sheet.cell_value(j, start[1])
    START = [7,4]

    hw = E/S
    lw = L/S
    print('h/w',hw)
    print('l/w',lw)
    if hw <= 1/2:
        pass
    elif hw <= 3/2:
        START[0]+=4
    elif hw < 6:
        START[0]+=8
    else:
        START[0]+=12

    if START[0] < 19:
        if lw <= 3/2:
            pass
        elif lw <4:
            START[0]+=2
    else:
        if lw == 3/2:
            pass
        elif lw == 1:
            START[0] += 2
        else:
            START[0] += 4
        
    ABCD0 = []
    ABCD90 = []
    for i in range(4):
        ABCD0.append(float(sheet.cell_value(START[0], START[1]+i)))
        ABCD90.append(float(sheet.cell_value(START[0]+1, START[1]+i)))

    print(ABCD0)
    print(ABCD90)

    WW1 =round((ABCD0[0]),2)
    WL1 =round((ABCD0[1]),2)
    GW1B =round((ABCD0[2]),2)
    GW1F =round((ABCD0[3]),2)

    WW3 =round((ABCD90[0]),2)
    WL3 =round((ABCD90[1]),2)
    GW3B =round((ABCD90[2]),2)
    GW3F =round((ABCD90[3]),2)


    loc = ("./table 6666.xls")
    wb = xlrd.open_workbook(loc)
    sheet = wb.sheet_by_index(0) # 

    # sheet.cell_value(j, start[1])
    START = [9,3]

    hw = E/S
    print('h/w',hw)

    if hw <= 1/2:
        pass
    elif hw <= 3/2:
        START[0]+=7
    elif hw < 6:
        START[0]+=14
    print(START)

    alpha_dc = {0:0,5:1,10:2,20:3,30:4,45:5,60:6}
    for k in alpha_dc.keys():
        if A == 0:
            break
        if A==60:
            START[0]+=5
            break
        if A <= k:
            START[0]+=alpha_dc[k]-1
            break
    print(START)
    ABCD0 = []
    ABCD90 = []
    for i in range(4):
        ABCD0.append(float(sheet.cell_value(START[0], START[1]+i)))
        ABCD90.append(float(sheet.cell_value(START[0]+1, START[1]+i)))
    print(ABCD0)
    print(ABCD90)


    XX = float(sheet.cell_value(START[0], START[1]-1))
    XY = float(sheet.cell_value(START[0]+1, START[1]-1))
    print(XX,XY)

    Z1 = ((A-XX) * ((ABCD90[0])-(ABCD0[0])))/((XY-XX))
    Z2 = ((A-XX) * ((ABCD90[1])-(ABCD0[1])))/((XY-XX))
    Z3 = ((A-XX) * ((ABCD90[2])-(ABCD0[2])))/((XY-XX))
    Z4 = ((A-XX) * ((ABCD90[3])-(ABCD0[3])))/((XY-XX))

    print(Z1)
    print(Z2)

    WR1 =round((Z1 + ABCD0[0]),2)
    LR1 =round((Z2 + ABCD0[1]),2)
    WR3 =round((Z3 + ABCD0[2]),2)
    LR3 =round((Z4 + ABCD0[3]),2)

    #*****************************************  case-----1-----  ****************************************
    document.add_paragraph('''Wind Load (F):''','heading 1')
    document.add_paragraph('''Case-1: 0 deg. internal suction''')
    document.add_paragraph(f'''Internal Wind Pressure                          = {IWP}
    External wind pressure for windward side wall   = {WW1} 
    External wind pressure for leeward side wall    = {WL1} 
    External wind pressure for windward side roof   = {WR1} 
    External wind pressure for leeward side roof    = {LR1} 
    External wind pressure for Front side Gabble wall    = {GW1F} 
    External wind pressure for Front side Gabble wall    = {GW1B} ''')

    C1WW = round(((IWPPLUS+WW1)*B*Pz),2)
    C1WL = round(((IWPMINUS+WL1)*B*Pz),2)
    C1RW = round(((IWPMINUS+WR1)*B*Pz),2)
    C1RL = round(((IWPMINUS+LR1)*B*Pz),2)

    table=document.add_table(rows=5,cols=7)
    table.style = "Table Grid"
    row = table.rows[0].cells
    row[0].text = "No."
    row[1].text = "Bay Lenght"
    row[2].text = "Wall Windward"
    row[3].text = "Wall Leeward"
    row[4].text = "Roof Windward"
    row[5].text = "Roof Leeward"

    row1 = table.rows[1].cells
    row1[0].text = "1"
    row1[1].text = (f"{B}")
    row1[2].text = (f"{C1WW}")
    row1[3].text = (f"{C1WL}")
    row1[4].text = (f"{C1RW}")
    row1[5].text = (f"{C1RL}")


    #*****************************************  case-----2-----  ****************************************

    document.add_paragraph('''Case-2: 0 deg. internal Pressure''')
    document.add_paragraph(f'''Internal Wind Pressure                          = {IWP}
    External wind pressure for windward side wall   = {WW1}
    External wind pressure for leeward side wall    = {WL1}
    External wind pressure for windward side roof   = {WR1}
    External wind pressure for leeward side roof    = {LR1}
    External wind pressure for Front side Gabble wall    = {GW1F}
    External wind pressure for Front side Gabble wall    = {GW1B} ''')

    C2WW = round(((IWPMINUS+WW1)*B*Pz),2)
    C2WL = round(((IWPPLUS+WL1)*B*Pz),2)
    C2RW = round(((IWPPLUS+WR1)*B*Pz),2)
    C2RL = round(((IWPPLUS+LR1)*B*Pz),2)

    table=document.add_table(rows=5,cols=7)
    table.style = "Table Grid"
    row = table.rows[0].cells
    row[0].text = "No."
    row[1].text = "Bay Lenght"
    row[2].text = "Wall Windward"
    row[3].text = "Wall Leeward"
    row[4].text = "Roof Windward"
    row[5].text = "Roof Leeward"

    row1 = table.rows[1].cells
    row1[0].text = "1"
    row1[1].text = (f"{B}")
    row1[2].text = (f"{C2WW}")
    row1[3].text = (f"{C2WL}")
    row1[4].text = (f"{C2RW}")
    row1[5].text = (f"{C2RL}")

    document.add_paragraph('''For gabble  wall''')

    table=document.add_table(rows=5,cols=4)
    table.style = "Table Grid"
    row = table.rows[0].cells
    row[0].text = "No."
    row[1].text = "Bay Lenght"
    row[2].text = "Front Gabble"
    row[3].text = "Back Gabble"

    row1 = table.rows[1].cells
    row1[0].text = "1"
    row1[1].text = (f"{B}")
    row1[2].text = (f"{C2WW}")
    row1[3].text = (f"{C2WL}")

    #*****************************************  case-----3-----  ****************************************

    document.add_paragraph('''Case-3: 90 deg. internal Pressure''')
    document.add_paragraph(f'''Internal Wind Pressure                          = {IWP}
    External wind pressure for windward side wall   = {WW3}  
    External wind pressure for leeward side wall    = {WL3} 
    External wind pressure for windward side roof   = {WR3} 
    External wind pressure for leeward side roof    = {LR3} 
    External wind pressure for Front side Gabble wall    = {GW3F} 
    External wind pressure for Front side Gabble wall    = {GW3B} ''')

    C2WW = round(((IWPMINUS-WW3)*B*Pz),2)
    C2WL = round(((IWPPLUS+WL3)*B*Pz),2)
    C2RW = round(((IWPMINUS-WR3)*B*Pz),2)
    C2RL = round(((IWPMINUS-LR3)*B*Pz),2)

    table=document.add_table(rows=5,cols=7)
    table.style = "Table Grid"
    row = table.rows[0].cells
    row[0].text = "No."
    row[1].text = "Bay Lenght"
    row[2].text = "Wall Windward"
    row[3].text = "Wall Leeward"
    row[4].text = "Roof Windward"
    row[5].text = "Roof Leeward"

    row1 = table.rows[1].cells
    row1[0].text = "1"
    row1[1].text = (f"{B}")
    row1[2].text = (f"{C2WW}")
    row1[3].text = (f"{C2WL}")
    row1[4].text = (f"{C2RW}")
    row1[5].text = (f"{C2RL}")







    document.add_paragraph('''STAAD PLANE
    START JOB INFORMATION
    ENGINEER DATE 04-Jun-21
    END JOB INFORMATION
    INPUT WIDTH 79
    UNIT METER KN''')


    document.add_paragraph(f'''
    JOINT COORDINATES
    1 0 0 0; 2 0 {E} 0; 3 {S1} {R} 0; 4 {S} {E} 0; 5 {S} 0 0;
    MEMBER INCIDENCES
    1 1 2; 2 2 3; 3 3 4; 4 4 5;
    SUPPORTS
    1 5 PINNED''')

    document.add_paragraph(f'''
    #*************************************
    #              DEAD LOAD
    #*************************************
    LOAD 1 LOADTYPE Dead  TITLE DEAD LOAD
    SELFWEIGHT Y -1 LIST ALL
    MEMBER LOAD
    2 3 UNI GY -{TF}
    #*************************************
    #              LIVE LOAD
    #*************************************
    LOAD 2 LOADTYPE Live  TITLE LIVE LOAD
    MEMBER LOAD
    2 3 UNI GY -{LLR}
    #*************************************
    #              WIND LOAD
    #*************************************
    #  *************** 0 DEG. WIND INT. SUCTION ********************
    LOAD 3 LOADTYPE Wind  TITLE 0 DEG. WIND INT. SUCTION
    MEMBER LOAD
    1 UNI GX 1
    2 UNI GY 2
    3 UNI GY 3
    4 UNI GX 4
    #  *************** 90 DEG. WIND INT. SUCTION ********************
    LOAD 4 LOADTYPE Wind  TITLE 90 DEG. WIND INT. SUCTION
    MEMBER LOAD
    1 UNI GX 1
    2 UNI GY 2
    3 UNI GY 3
    4 UNI GX 4
    #  *************** 90 DEG. WIND INT. PRESSURE ********************
    LOAD 5 LOADTYPE Wind  TITLE 90 DEG. WIND INT. PRESSURE
    MEMBER LOAD
    1 UNI GX 1
    2 UNI GY 2
    3 UNI GY 3
    4 UNI GX 4
    PERFORM ANALYSIS PRINT ALL
    PARAMETER 1
    CODE IS800 LSD
    CHECK CODE ALL
    FINISH''')





    docx_filename = "WL.docx"
    document.save("WL.docx")
    
    pdf_filename = "WL.pdf"
    convert(docx_filename, pdf_filename)

    # Send the PDF file as a downloadable response
    return send_file(pdf_filename, as_attachment=True)
    # Return the generated document as a downloadable file
    # return send_file('WL.docx', as_attachment=True)
    os.remove(docx_filename)
    os.remove(pdf_filename)

if __name__ == '__main__':
    app.run(debug=True)
